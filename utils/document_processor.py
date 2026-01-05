"""
Document Processing Pipeline for Q&A System
Handles file type detection, text extraction, and intelligent chunking.
"""

import os
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import mimetypes
from dataclasses import dataclass
from abc import ABC, abstractmethod

# Import document processing libraries
try:
    import fitz  # PyMuPDF for PDF processing
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document  # python-docx for Word documents
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of document content."""
    content: str
    chunk_id: str
    document_id: str
    page_number: Optional[int] = None
    section_title: Optional[str] = None
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class DocumentMetadata:
    """Metadata for a processed document."""
    document_id: str
    filename: str
    file_type: str
    file_size: int
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    creation_date: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None


class DocumentProcessor(ABC):
    """Abstract base class for document processors."""

    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the file type."""
        pass

    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        """Extract text content from the document."""
        pass

    @abstractmethod
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract metadata from the document."""
        pass


class PDFProcessor(DocumentProcessor):
    """PDF document processor using PyMuPDF."""

    def can_process(self, file_path: Path) -> bool:
        if not PYMUPDF_AVAILABLE:
            return False
        mime_type, _ = mimetypes.guess_type(str(file_path))
        return mime_type == 'application/pdf'

    def extract_text(self, file_path: Path) -> str:
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available for PDF processing")

        text = ""
        with fitz.open(str(file_path)) as doc:
            for page in doc:
                text += page.get_text() + "\n"

        return text.strip()

    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available for PDF processing")

        with fitz.open(str(file_path)) as doc:
            metadata = doc.metadata

            return DocumentMetadata(
                document_id=str(file_path.stem),
                filename=file_path.name,
                file_type="pdf",
                file_size=file_path.stat().st_size,
                page_count=len(doc),
                creation_date=metadata.get("creationDate"),
                author=metadata.get("author"),
                title=metadata.get("title"),
            )


class DOCXProcessor(DocumentProcessor):
    """Word document processor using python-docx."""

    def can_process(self, file_path: Path) -> bool:
        if not DOCX_AVAILABLE:
            return False
        mime_type, _ = mimetypes.guess_type(str(file_path))
        return mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]

    def extract_text(self, file_path: Path) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")

        doc = Document(str(file_path))
        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text.strip()

    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")

        doc = Document(str(file_path))
        core_props = doc.core_properties

        return DocumentMetadata(
            document_id=str(file_path.stem),
            filename=file_path.name,
            file_type="docx",
            file_size=file_path.stat().st_size,
            creation_date=str(core_props.created) if core_props.created else None,
            author=core_props.author,
            title=core_props.title,
        )


class TextProcessor(DocumentProcessor):
    """Plain text file processor."""

    def can_process(self, file_path: Path) -> bool:
        mime_type, _ = mimetypes.guess_type(str(file_path))
        return mime_type and mime_type.startswith('text/')

    def extract_text(self, file_path: Path) -> str:
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, raise error
            raise UnicodeDecodeError("Could not decode file with any supported encoding")

        except Exception as e:
            raise ValueError(f"Failed to read text file: {e}")

    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        text = self.extract_text(file_path)

        return DocumentMetadata(
            document_id=str(file_path.stem),
            filename=file_path.name,
            file_type="text",
            file_size=file_path.stat().st_size,
            word_count=len(text.split()),
            character_count=len(text),
        )


class DocumentProcessingPipeline:
    """Main document processing pipeline."""

    def __init__(self):
        self.processors: List[DocumentProcessor] = [
            PDFProcessor(),
            DOCXProcessor(),
            TextProcessor(),
        ]

        # Log available processors
        available = []
        for processor in self.processors:
            if hasattr(processor, '__class__'):
                name = processor.__class__.__name__.replace('Processor', '')
                if processor.can_process(Path("dummy.pdf")) or name == "Text":  # Rough check
                    available.append(name)
        logger.info(f"Available document processors: {', '.join(available)}")

    def validate_file(self, file_path: Path, max_size_mb: int = 10) -> bool:
        """Validate file before processing."""
        if not file_path.exists():
            raise FileNotFoundError(f"File does not exist: {file_path}")

        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")

        # Check file size
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_path.stat().st_size > max_size_bytes:
            raise ValueError(f"File size ({file_path.stat().st_size / (1024*1024):.1f}MB) exceeds maximum ({max_size_mb}MB)")

        # Check if we have a processor for this file type
        processor = self._get_processor(file_path)
        if not processor:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        return True

    def _get_processor(self, file_path: Path) -> Optional[DocumentProcessor]:
        """Get the appropriate processor for a file."""
        for processor in self.processors:
            try:
                if processor.can_process(file_path):
                    return processor
            except Exception:
                continue
        return None

    def process_document(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Process a document and return text content and metadata."""
        self.validate_file(file_path)

        processor = self._get_processor(file_path)
        if not processor:
            raise ValueError(f"No processor available for {file_path}")

        try:
            text = processor.extract_text(file_path)
            metadata = processor.extract_metadata(file_path)

            logger.info(f"Processed document: {file_path.name} ({len(text)} characters)")
            return text, metadata

        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {e}")
            raise

    def chunk_text(
        self,
        text: str,
        document_id: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        strategy: str = "sentence"
    ) -> List[DocumentChunk]:
        """Chunk text into manageable pieces."""

        if strategy == "sentence":
            return self._chunk_by_sentences(text, document_id, chunk_size, chunk_overlap)
        elif strategy == "paragraph":
            return self._chunk_by_paragraphs(text, document_id, chunk_size, chunk_overlap)
        elif strategy == "fixed":
            return self._chunk_fixed_size(text, document_id, chunk_size, chunk_overlap)
        else:
            raise ValueError(f"Unknown chunking strategy: {strategy}")

    def _chunk_by_sentences(
        self,
        text: str,
        document_id: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> List[DocumentChunk]:
        """Chunk text by sentences."""
        import re

        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text.strip())

        chunks = []
        current_chunk = ""
        chunk_id = 0

        for sentence in sentences:
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                # Create chunk
                chunk = DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_id=f"{document_id}_chunk_{chunk_id}",
                    document_id=document_id,
                )
                chunks.append(chunk)

                # Start new chunk with overlap
                overlap_text = current_chunk[-chunk_overlap:] if chunk_overlap > 0 else ""
                current_chunk = overlap_text + " " + sentence
                chunk_id += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence

        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                chunk_id=f"{document_id}_chunk_{chunk_id}",
                document_id=document_id,
            )
            chunks.append(chunk)

        return chunks

    def _chunk_by_paragraphs(
        self,
        text: str,
        document_id: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> List[DocumentChunk]:
        """Chunk text by paragraphs."""
        paragraphs = text.split('\n\n')

        chunks = []
        current_chunk = ""
        chunk_id = 0

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                # Create chunk
                chunk = DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_id=f"{document_id}_chunk_{chunk_id}",
                    document_id=document_id,
                )
                chunks.append(chunk)

                # Start new chunk with overlap
                overlap_text = current_chunk[-chunk_overlap:] if chunk_overlap > 0 else ""
                current_chunk = overlap_text + "\n\n" + paragraph
                chunk_id += 1
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph

        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                chunk_id=f"{document_id}_chunk_{chunk_id}",
                document_id=document_id,
            )
            chunks.append(chunk)

        return chunks

    def _chunk_fixed_size(
        self,
        text: str,
        document_id: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> List[DocumentChunk]:
        """Chunk text into fixed-size pieces."""
        chunks = []
        start = 0
        chunk_id = 0

        while start < len(text):
            end = min(start + chunk_size, len(text))

            # If we're not at the end, try to break at a word boundary
            if end < len(text):
                # Look for last space within overlap distance from end
                last_space = text.rfind(' ', start, end)
                if last_space > start + chunk_size - chunk_overlap:
                    end = last_space

            chunk_content = text[start:end].strip()
            if chunk_content:
                chunk = DocumentChunk(
                    content=chunk_content,
                    chunk_id=f"{document_id}_chunk_{chunk_id}",
                    document_id=document_id,
                )
                chunks.append(chunk)
                chunk_id += 1

            # Move start position with overlap
            start = end - chunk_overlap if chunk_overlap > 0 else end

        return chunks